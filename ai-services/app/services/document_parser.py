# app/services/document_parser.py
"""
文档解析服务
支持 DOCX / PDF / 图片（OCR）格式解析
保留页码、段落索引和位置信息
"""

import logging
import os
import re
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class ParagraphInfo:
    """段落信息"""
    index: int                    # 段落序号（从1开始）
    text: str                     # 段落文本
    page: Optional[int] = None    # 页码（从1开始）
    bbox: Optional[Tuple[float, float, float, float]] = None  # 位置边界框


@dataclass
class ParseResult:
    """解析结果"""
    full_text: str                          # 完整合同文本
    paragraphs: List[ParagraphInfo] = field(default_factory=list)
    page_count: int = 0                     # 总页数
    format_type: str = ""                   # 原始格式：docx / pdf / image
    warnings: List[str] = field(default_factory=list)  # 解析告警
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "full_text": self.full_text,
            "paragraphs": [
                {
                    "index": p.index,
                    "text": p.text,
                    "page": p.page,
                    "bbox": list(p.bbox) if p.bbox else None
                }
                for p in self.paragraphs
            ],
            "page_count": self.page_count,
            "format_type": self.format_type,
            "warnings": self.warnings,
            "metadata": self.metadata
        }


# ============================================================
# 文件格式检测
# ============================================================

def detect_format(file_path: str) -> str:
    """
    检测文件格式
    返回: "docx" | "pdf" | "image" | "txt"
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.docx':
        return 'docx'
    elif ext == '.pdf':
        return 'pdf'
    elif ext in ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif', '.gif', '.webp'):
        return 'image'
    elif ext == '.txt':
        return 'txt'
    else:
        # 尝试通过MIME判断
        raise ValueError(f"不支持的文件格式: {ext}")


# ============================================================
# DOCX 解析
# ============================================================

def parse_docx(file_path: str) -> ParseResult:
    """
    解析 DOCX 文件
    提取文本和段落结构，保留段落索引
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    logger.info(f"开始解析 DOCX 文件: {file_path}")
    warnings = []

    try:
        doc = Document(file_path)

        paragraphs = []
        full_text_parts = []
        para_index = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            para_index += 1
            paragraphs.append(ParagraphInfo(
                index=para_index,
                text=text,
                page=None  # python-docx 不直接提供页码
            ))
            full_text_parts.append(text)

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    para_index += 1
                    combined = " | ".join(row_texts)
                    paragraphs.append(ParagraphInfo(
                        index=para_index,
                        text=combined,
                        page=None
                    ))
                    full_text_parts.append(combined)

        full_text = "\n\n".join(full_text_parts)

        result = ParseResult(
            full_text=full_text,
            paragraphs=paragraphs,
            page_count=1,  # DOCX 没有固定分页
            format_type="docx",
            warnings=warnings,
            metadata={
                "paragraph_count": para_index,
                "has_tables": len(doc.tables) > 0
            }
        )

        logger.info(f"DOCX 解析完成: {para_index} 个段落, {len(full_text)} 字符")
        return result

    except Exception as e:
        logger.error(f"DOCX 解析失败: {e}")
        raise


# ============================================================
# PDF 解析 (pdfplumber 为主)
# ============================================================

def parse_pdf(file_path: str) -> ParseResult:
    """
    解析 PDF 文件
    使用 pdfplumber 提取文本，保留页码信息
    失败时尝试 PyPDF2 作为备用方案
    """
    logger.info(f"开始解析 PDF 文件: {file_path}")
    warnings = []

    # 优先使用 pdfplumber
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("请安装 pdfplumber: pip install pdfplumber")

    try:
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            paragraphs = []
            full_text_parts = []
            para_index = 0

            for page_num, page in enumerate(pdf.pages, start=1):
                # 提取页面文本
                page_text = page.extract_text()
                if not page_text:
                    warnings.append(f"第{page_num}页无法提取文本（可能是扫描件或图片型PDF）")
                    continue

                # 按段落分割
                page_paragraphs = page_text.split('\n\n')
                for para_text in page_paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue

                    # 清理换行符（PDF 提取常有多余换行）
                    para_text = _clean_pdf_paragraph(para_text)
                    if not para_text:
                        continue

                    para_index += 1
                    paragraphs.append(ParagraphInfo(
                        index=para_index,
                        text=para_text,
                        page=page_num
                    ))
                    full_text_parts.append(para_text)

            full_text = "\n\n".join(full_text_parts)

            # 如果没有提取到文本，尝试 PyPDF2
            if not full_text.strip():
                warnings.append("pdfplumber 未提取到文本，尝试 PyPDF2 备用方案")
                return _parse_pdf_fallback(file_path, warnings)

        result = ParseResult(
            full_text=full_text,
            paragraphs=paragraphs,
            page_count=page_count,
            format_type="pdf",
            warnings=warnings,
            metadata={
                "paragraph_count": para_index,
                "parser": "pdfplumber"
            }
        )

        logger.info(f"PDF 解析完成: {page_count} 页, {para_index} 个段落")
        return result

    except Exception as e:
        logger.warning(f"pdfplumber 解析失败: {e}，尝试备用方案")
        return _parse_pdf_fallback(file_path, warnings + [str(e)])


def _parse_pdf_fallback(file_path: str, warnings: List[str]) -> ParseResult:
    """PDF 解析备用方案：使用 PyPDF2"""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("请安装 PyPDF2: pip install PyPDF2")

    reader = PdfReader(file_path)
    page_count = len(reader.pages)
    paragraphs = []
    full_text_parts = []
    para_index = 0

    for page_num, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text()
        if not page_text:
            warnings.append(f"第{page_num}页无法提取文本")
            continue

        page_paragraphs = page_text.split('\n\n')
        for para_text in page_paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            para_text = _clean_pdf_paragraph(para_text)
            if not para_text:
                continue

            para_index += 1
            paragraphs.append(ParagraphInfo(
                index=para_index,
                text=para_text,
                page=page_num
            ))
            full_text_parts.append(para_text)

    full_text = "\n\n".join(full_text_parts)

    return ParseResult(
        full_text=full_text,
        paragraphs=paragraphs,
        page_count=page_count,
        format_type="pdf",
        warnings=warnings,
        metadata={
            "paragraph_count": para_index,
            "parser": "PyPDF2"
        }
    )


def _clean_pdf_paragraph(text: str) -> str:
    """清理 PDF 提取的段落文本"""
    # 去掉行内多余空格
    text = re.sub(r' +', ' ', text)
    # 去掉行首行尾空白
    text = text.strip()
    # 合并被断开的行（单行换行符替换为空格）
    lines = text.split('\n')
    if len(lines) > 1:
        # 检查是否是同一段落的续行
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
        text = ' '.join(cleaned_lines)
    return text


# ============================================================
# 图片 OCR 解析
# ============================================================

def parse_image(file_path: str, use_paddleocr: bool = True) -> ParseResult:
    """
    解析图片文件（OCR 识别）

    Args:
        file_path: 图片文件路径
        use_paddleocr: 是否优先使用 PaddleOCR（默认True，不可用时降级）

    Returns:
        ParseResult
    """
    logger.info(f"开始 OCR 解析图片: {file_path}")
    warnings = []

    full_text = ""
    paragraphs = []

    if use_paddleocr:
        try:
            full_text, paragraphs = _ocr_with_paddleocr(file_path)
            if full_text.strip():
                return ParseResult(
                    full_text=full_text,
                    paragraphs=paragraphs,
                    page_count=1,
                    format_type="image",
                    warnings=warnings,
                    metadata={"ocr_engine": "PaddleOCR"}
                )
        except ImportError:
            warnings.append("PaddleOCR 未安装，使用基础图像处理")
        except Exception as e:
            warnings.append(f"PaddleOCR 识别失败: {str(e)}")

    # 降级：尝试基础图像信息提取
    try:
        full_text, paragraphs = _ocr_fallback(file_path)
        warnings.append("图片OCR未识别到文本，可能需要安装PaddleOCR")
    except Exception as e:
        warnings.append(f"图片解析失败: {str(e)}")

    return ParseResult(
        full_text=full_text,
        paragraphs=paragraphs,
        page_count=1,
        format_type="image",
        warnings=warnings,
        metadata={"ocr_engine": "fallback"}
    )


def _ocr_with_paddleocr(file_path: str) -> Tuple[str, List[ParagraphInfo]]:
    """使用 PaddleOCR 进行文字识别"""
    from paddleocr import PaddleOCR

    ocr = PaddleOCR(use_angle_cls=True, lang='ch', show_log=False)
    result = ocr.ocr(file_path, cls=True)

    if not result or not result[0]:
        return "", []

    paragraphs = []
    full_text_parts = []
    para_index = 0

    for line_info in result[0]:
        bbox = line_info[0]  # [[x1,y1],[x2,y2],[x3,y3],[x4,y4]]
        text = line_info[1][0]  # 识别文本
        confidence = line_info[1][1]  # 置信度

        if confidence < 0.5:  # 低置信度过滤
            continue

        para_index += 1
        # 计算边界框 (x1, y1, x2, y2)
        x_coords = [p[0] for p in bbox]
        y_coords = [p[1] for p in bbox]
        bbox_tuple = (min(x_coords), min(y_coords), max(x_coords), max(y_coords))

        paragraphs.append(ParagraphInfo(
            index=para_index,
            text=text,
            page=1,
            bbox=bbox_tuple
        ))
        full_text_parts.append(text)

    return "\n".join(full_text_parts), paragraphs


def _ocr_fallback(file_path: str) -> Tuple[str, List[ParagraphInfo]]:
    """OCR 降级方案：提取图片基本信息"""
    from PIL import Image

    img = Image.open(file_path)
    width, height = img.size

    metadata_text = (
        f"[图片信息] 文件名: {os.path.basename(file_path)}, "
        f"尺寸: {width}x{height}, "
        f"模式: {img.mode}"
    )

    # 如果图片中有嵌入的文本元数据，尝试读取
    exif_text = ""
    try:
        from PIL.ExifTags import TAGS
        exif = img._getexif()
        if exif:
            for tag_id, value in exif.items():
                tag_name = TAGS.get(tag_id, tag_id)
                if isinstance(value, str) and len(value) > 10:
                    exif_text += f"{value}\n"
    except Exception:
        pass

    text = exif_text if exif_text.strip() else metadata_text
    return text, []


# ============================================================
# 统一解析入口
# ============================================================

def parse_document(
    file_path: str,
    format_type: Optional[str] = None,
    use_ocr: bool = True
) -> ParseResult:
    """
    统一文档解析入口

    Args:
        file_path: 文档文件路径
        format_type: 文件格式（可选，自动检测）
        use_ocr: 是否使用OCR（仅图片格式）

    Returns:
        ParseResult: 解析结果
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    file_size = os.path.getsize(file_path)
    if file_size == 0:
        raise ValueError("文件为空")

    if format_type is None:
        format_type = detect_format(file_path)

    logger.info(f"解析文档: {file_path}, 格式: {format_type}, 大小: {file_size} bytes")

    if format_type == 'docx':
        result = parse_docx(file_path)
    elif format_type == 'pdf':
        result = parse_pdf(file_path)
    elif format_type == 'image':
        result = parse_image(file_path, use_paddleocr=use_ocr)
    elif format_type == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        result = ParseResult(
            full_text=text,
            paragraphs=[
                ParagraphInfo(index=i+1, text=line)
                for i, line in enumerate(lines)
            ],
            page_count=1,
            format_type="txt",
            metadata={"paragraph_count": len(lines)}
        )
    else:
        raise ValueError(f"不支持的格式: {format_type}")

    return result


def parse_document_bytes(
    content: bytes,
    filename: str,
    use_ocr: bool = True
) -> ParseResult:
    """
    解析字节形式的文档（用于API上传）

    Args:
        content: 文件字节内容
        filename: 原始文件名（用于判断格式）
        use_ocr: 是否使用OCR

    Returns:
        ParseResult
    """
    import tempfile

    ext = os.path.splitext(filename)[1].lower()
    suffix = ext if ext else '.tmp'

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        result = parse_document(tmp_path, use_ocr=use_ocr)
        result.metadata['original_filename'] = filename
        return result
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
