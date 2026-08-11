"""文档解析器 - 支持DOCX、PDF格式，提取文本、段落和页码信息"""
import logging
import os
from typing import Dict, Any, List, Tuple

from app.core.exceptions import FileParseError
from app.utils.text_utils import clean_text

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器，支持DOCX/PDF文件解析，输出结构化文本段"""

    SUPPORTED_TYPES = {"docx", "pdf", "image"}

    def parse(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        解析文档

        Args:
            file_path: 文件路径
            file_type: 文件类型 (docx/pdf/image)

        Returns:
            {"text": "全文文本", "segments": [...], "warnings": [...]}
        """
        if not os.path.exists(file_path):
            raise FileParseError(f"文件不存在: {file_path}")

        file_type = file_type.lower().strip()
        if file_type not in self.SUPPORTED_TYPES:
            raise FileParseError(
                f"不支持的文件类型: {file_type}",
                detail={"supported_types": list(self.SUPPORTED_TYPES)},
            )

        if file_type == "docx":
            return self._parse_docx(file_path)
        elif file_type == "pdf":
            return self._parse_pdf(file_path)
        elif file_type == "image":
            # 图片类型委托给OCR处理器
            return self._parse_image_placeholder(file_path)

        raise FileParseError(f"未处理的文件类型: {file_type}")

    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析DOCX文件"""
        try:
            from docx import Document

            doc = Document(file_path)
            segments = []
            all_text_parts = []
            para_idx = 0

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    segments.append({
                        "text": text,
                        "page": None,  # python-docx不直接提供页码
                        "paragraph_index": para_idx,
                        "style": paragraph.style.name if paragraph.style else "Normal",
                    })
                    all_text_parts.append(text)
                    para_idx += 1

            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        combined = " | ".join(row_texts)
                        segments.append({
                            "text": combined,
                            "page": None,
                            "paragraph_index": para_idx,
                            "style": "TableRow",
                        })
                        all_text_parts.append(combined)
                        para_idx += 1

            full_text = "\n\n".join(all_text_parts)
            full_text = clean_text(full_text)

            warnings = self._check_docx_warnings(doc, full_text)

            return {
                "text": full_text,
                "segments": segments,
                "warnings": warnings,
            }

        except ImportError:
            raise FileParseError(
                "缺少python-docx依赖",
                detail={"hint": "请安装: pip install python-docx"},
            )
        except Exception as e:
            logger.error(f"DOCX解析失败: {str(e)}")
            raise FileParseError(f"DOCX文档解析失败: {str(e)}")

    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析PDF文件"""
        try:
            import pdfplumber

            segments = []
            all_text_parts = []
            page_count = 0
            para_global = 0

            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_count = page_num
                    text = page.extract_text()
                    if text:
                        paragraphs = text.split("\n\n") if "\n\n" in text else text.split("\n")
                        for para_text in paragraphs:
                            para_text = para_text.strip()
                            if para_text:
                                segments.append({
                                    "text": para_text,
                                    "page": page_num,
                                    "paragraph_index": para_global,
                                })
                                all_text_parts.append(para_text)
                                para_global += 1

                    # 提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            row_texts = [c.strip() if c else "" for c in row]
                            row_text = " | ".join(filter(None, row_texts))
                            if row_text.strip():
                                segments.append({
                                    "text": row_text,
                                    "page": page_num,
                                    "paragraph_index": para_global,
                                    "style": "TableRow",
                                })
                                all_text_parts.append(row_text)
                                para_global += 1

            full_text = "\n\n".join(all_text_parts)
            full_text = clean_text(full_text)

            warnings = self._check_pdf_warnings(page_count, full_text)

            return {
                "text": full_text,
                "segments": segments,
                "warnings": warnings,
            }

        except ImportError:
            raise FileParseError(
                "缺少pdfplumber依赖",
                detail={"hint": "请安装: pip install pdfplumber"},
            )
        except Exception as e:
            logger.error(f"PDF解析失败: {str(e)}")
            raise FileParseError(f"PDF文档解析失败: {str(e)}")

    def _parse_image_placeholder(self, file_path: str) -> Dict[str, Any]:
        """图片文件占位，实际OCR由OCRProcessor处理"""
        return {
            "text": "",
            "segments": [],
            "warnings": ["图片文件需经过OCR识别，请调用OCR接口"],
        }

    @staticmethod
    def _check_docx_warnings(doc, full_text: str) -> List[str]:
        """检查DOCX文档质量"""
        warnings = []
        if len(full_text) < 50:
            warnings.append("文档文本内容过少，可能为空文档或扫描件")
        if doc.tables and len(doc.tables) > 20:
            warnings.append("文档包含大量表格，部分内容可能为表格格式")
        return warnings

    @staticmethod
    def _check_pdf_warnings(page_count: int, full_text: str) -> List[str]:
        """检查PDF文档质量"""
        warnings = []
        if page_count == 0:
            warnings.append("PDF文档为空白")
        if len(full_text) < 50:
            warnings.append("PDF文本内容过少，可能为扫描件，建议使用OCR")
        return warnings
